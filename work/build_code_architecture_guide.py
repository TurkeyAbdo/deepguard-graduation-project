from __future__ import annotations

import math
import re
import sys
from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
APP_ROOT = ROOT / "deepguard"
OUT_DIR = ROOT / "outputs"
ASSET_DIR = ROOT / "work" / "code_architecture_assets"
OUTPUT = OUT_DIR / "DEEPGUARD_FULL_CODE_ARCHITECTURE_GUIDE.docx"

NAVY = "#16324F"
BLUE = "#2E74B5"
DARK_BLUE = "#1F4D78"
TEAL = "#167D83"
GREEN = "#2C7A5A"
AMBER = "#9A6A16"
RED = "#A23A3A"
INK = "#20262E"
MUTED = "#5B6673"
LINE = "#CBD4DE"
PALE_BLUE = "#E8EEF5"
PALE_TEAL = "#E7F2F1"
PALE_GREEN = "#E8F3ED"
PALE_AMBER = "#F7F0DE"
PALE_RED = "#F8E9E8"
LIGHT = "#F4F6F9"
WHITE = "#FFFFFF"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_SIDE_DXA = 120


SOURCE_GROUPS: list[tuple[str, list[str]]] = [
    (
        "A. Application Entry, Shared Types, and Styling",
        [
            "deepguard/package.json",
            "deepguard/requirements.txt",
            "deepguard/run.ps1",
            "deepguard/vite.config.ts",
            "deepguard/tsconfig.json",
            "deepguard/tsconfig.app.json",
            "deepguard/tsconfig.node.json",
            "deepguard/.oxlintrc.json",
            "deepguard/src/main.tsx",
            "deepguard/src/App.tsx",
            "deepguard/src/types.ts",
            "deepguard/src/index.css",
            "deepguard/src/App.css",
            "deepguard/src/components/StatusBadge.tsx",
        ],
    ),
    (
        "B. Verification, Browser AI, and Decision Logic",
        [
            "deepguard/src/views/VerificationView.tsx",
            "deepguard/src/lib/face.ts",
            "deepguard/src/lib/liveness.ts",
            "deepguard/src/lib/deepfake.ts",
            "deepguard/src/lib/decision.ts",
            "deepguard/src/lib/api.ts",
        ],
    ),
    (
        "C. Sessions, Evaluation, and System Views",
        [
            "deepguard/src/views/SessionsView.tsx",
            "deepguard/src/views/EvaluationView.tsx",
            "deepguard/src/views/SystemView.tsx",
            "deepguard/src/data/evaluationCases.ts",
            "deepguard/src/lib/evaluation.ts",
        ],
    ),
    (
        "D. Backend and Automated Tests",
        [
            "deepguard/server/__init__.py",
            "deepguard/server/app.py",
            "deepguard/server/test_app.py",
            "deepguard/src/test/setup.ts",
            "deepguard/src/lib/decision.test.ts",
            "deepguard/src/lib/liveness.test.ts",
            "deepguard/src/lib/evaluation.test.ts",
        ],
    ),
    (
        "E. Demo, Evaluation, and Report Tooling",
        [
            "deepguard/scripts/copy-mediapipe-assets.mjs",
            "deepguard/scripts/create_liveness_steps_video.py",
            "deepguard/scripts/create_synthetic_attack_video.py",
            "deepguard/scripts/generate_evaluation_artifacts.py",
            "work/add_evaluation_to_report.py",
        ],
    ),
]


PURPOSES = {
    "deepguard/package.json": "Frontend dependencies, scripts, test commands, and MediaPipe post-install step.",
    "deepguard/requirements.txt": "Pinned Python backend and test dependencies.",
    "deepguard/run.ps1": "One-command local launcher for API and browser application.",
    "deepguard/vite.config.ts": "Vite, Vitest, React plugin, and /api development proxy configuration.",
    "deepguard/tsconfig.json": "TypeScript project references.",
    "deepguard/tsconfig.app.json": "Strict browser application compiler settings.",
    "deepguard/tsconfig.node.json": "Strict build-tool compiler settings.",
    "deepguard/.oxlintrc.json": "React and hooks lint rules.",
    "deepguard/src/main.tsx": "React entry point and StrictMode mount.",
    "deepguard/src/App.tsx": "Application shell, navigation, lazy-loaded views, and session refresh coordination.",
    "deepguard/src/types.ts": "Shared contracts for challenges, decisions, sessions, metrics, and source integrity.",
    "deepguard/src/index.css": "Production interface tokens, layout, responsive rules, and component styling.",
    "deepguard/src/App.css": "Unused Vite starter stylesheet retained as legacy residue.",
    "deepguard/src/components/StatusBadge.tsx": "Reusable semantic status label.",
    "deepguard/src/views/VerificationView.tsx": "Camera lifecycle, challenge orchestration, inference, result display, and persistence.",
    "deepguard/src/lib/face.ts": "MediaPipe initialization, landmark/blendshape extraction, quality, yaw, and frame capture.",
    "deepguard/src/lib/liveness.ts": "Blink, smile, and ordered head-motion state machines.",
    "deepguard/src/lib/deepfake.ts": "Transformers.js model loading and three-frame deepfake scoring.",
    "deepguard/src/lib/decision.ts": "Explainable threshold policy for genuine, fake, and manual review outcomes.",
    "deepguard/src/lib/api.ts": "Typed browser client for FastAPI endpoints.",
    "deepguard/src/views/SessionsView.tsx": "Session history, filters, trend chart, details, and review actions.",
    "deepguard/src/views/EvaluationView.tsx": "Controlled evaluation dashboard and resource evidence.",
    "deepguard/src/views/SystemView.tsx": "Runtime capability summary and lightweight architecture explanation.",
    "deepguard/src/data/evaluationCases.ts": "Controlled genuine and replay-attack evaluation cases.",
    "deepguard/src/lib/evaluation.ts": "Confusion matrix, coverage, accuracy, F1, and latency calculations.",
    "deepguard/server/__init__.py": "Backend package marker.",
    "deepguard/server/app.py": "FastAPI application, validation, SQLite schema, API routes, and static hosting.",
    "deepguard/server/test_app.py": "API health, persistence, metrics, and review-state tests.",
    "deepguard/src/test/setup.ts": "Vitest browser test setup.",
    "deepguard/src/lib/decision.test.ts": "Decision-policy boundary tests.",
    "deepguard/src/lib/liveness.test.ts": "Challenge stage and ordered head-motion tests.",
    "deepguard/src/lib/evaluation.test.ts": "Evaluation-metric correctness tests.",
    "deepguard/scripts/copy-mediapipe-assets.mjs": "Copies MediaPipe WASM files into the public build.",
    "deepguard/scripts/create_liveness_steps_video.py": "Builds the professor-facing synthetic liveness instruction video.",
    "deepguard/scripts/create_synthetic_attack_video.py": "Builds a safe synthetic replay/deepfake demonstration video.",
    "deepguard/scripts/generate_evaluation_artifacts.py": "Exports controlled cases, metrics, and report figures.",
    "work/add_evaluation_to_report.py": "Integrates evaluation evidence into the graduation report DOCX.",
}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color.lstrip("#"))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.lstrip("#"))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Sequence[int], indent: int = TABLE_INDENT_DXA) -> None:
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_DXA}: {widths}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(
                cell,
                CELL_TOP_BOTTOM_DXA,
                CELL_SIDE_DXA,
                CELL_TOP_BOTTOM_DXA,
                CELL_SIDE_DXA,
            )
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_table_borders(table, color: str = LINE, size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color.lstrip("#"))


def set_paragraph_border(paragraph, color=LINE, size=6, space=6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color.lstrip("#"))
    p_bdr.append(bottom)


def shade_paragraph(paragraph, fill: str, border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill.lstrip("#"))
    if border:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        for edge in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "3")
            el.set(qn("w:color"), border.lstrip("#"))
            p_bdr.append(el)


def keep_with_next(paragraph, value=True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def set_run_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr, fld_char_2])


def set_image_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = rgb(DARK_BLUE)
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False

    if "Code Listing" not in styles:
        styles.add_style("Code Listing", 1)
    code = styles["Code Listing"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(7.5)
    code.font.color.rgb = rgb("23272D")
    code.paragraph_format.left_indent = Inches(0.08)
    code.paragraph_format.right_indent = Inches(0.08)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(4)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.keep_together = False
    code.paragraph_format.widow_control = False

    if "Code File" not in styles:
        styles.add_style("Code File", 1)
    code_file = styles["Code File"]
    code_file.font.name = "Calibri"
    code_file._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    code_file._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    code_file.font.size = Pt(13)
    code_file.font.bold = True
    code_file.font.color.rgb = rgb(DARK_BLUE)
    code_file.paragraph_format.space_before = Pt(12)
    code_file.paragraph_format.space_after = Pt(5)
    code_file.paragraph_format.keep_with_next = True
    code_file.paragraph_format.page_break_before = False


def create_numbering(document: Document) -> tuple[int, int]:
    numbering = document.part.numbering_part.element
    existing_abstract = [
        int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]

    def make_abstract(abstract_id: int, num_fmt: str, text: str, font: str | None = None) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        level.append(fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        level.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        level.append(jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        level.append(p_pr)
        if font:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            r_pr.append(fonts)
            level.append(r_pr)
        abstract.append(level)
        numbering.append(abstract)

    next_abstract = max(existing_abstract, default=-1) + 1
    make_abstract(next_abstract, "bullet", "\u2022", "Calibri")
    make_abstract(next_abstract + 1, "decimal", "%1.")

    ids = []
    next_num = max(existing_num, default=0) + 1
    for offset, abstract_id in enumerate((next_abstract, next_abstract + 1)):
        num_id = next_num + offset
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num_id = OxmlElement("w:abstractNumId")
        abstract_num_id.set(qn("w:val"), str(abstract_id))
        num.append(abstract_num_id)
        numbering.append(num)
        ids.append(num_id)
    return ids[0], ids[1]


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def add_bullet(document: Document, text: str, bullet_id: int, bold_prefix: str | None = None):
    p = document.add_paragraph()
    apply_numbering(p, bullet_id)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2)
    else:
        set_run_font(p.add_run(text))
    return p


def add_numbered(document: Document, text: str, number_id: int):
    p = document.add_paragraph()
    apply_numbering(p, number_id)
    set_run_font(p.add_run(text))
    return p


def add_callout(
    document: Document,
    label: str,
    text: str,
    fill: str = PALE_BLUE,
    accent: str = BLUE,
):
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    shade_paragraph(p, fill, accent)
    r = p.add_run(f"{label.upper()}  ")
    set_run_font(r, size=9.5, color=accent, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    return p


def add_table(
    document: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths: Sequence[int],
    center_columns: Iterable[int] = (),
):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    center_columns = set(center_columns)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in center_columns else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        set_run_font(p.add_run(str(value)), size=9, color=NAVY, bold=True)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in center_columns else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            set_run_font(p.add_run(str(value)), size=8.5, color=INK)
    set_table_geometry(table, widths)
    set_table_borders(table)
    after = document.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)
    return table


def add_picture(document: Document, path: Path, width: float, title: str, alt: str, caption: str):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    set_image_alt_text(shape, title, alt)
    cap = document.add_paragraph(caption, style="Caption")
    return shape


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/segoeui.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    if bold:
        candidates = [
            Path("C:/Windows/Fonts/segoeuib.ttf"),
            Path("C:/Windows/Fonts/arialbd.ttf"),
        ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def rounded_box(draw, box, fill, outline, title, body="", title_color=INK):
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = box
    title_font = font(30, True)
    body_font = font(20)
    draw.text((x1 + 24, y1 + 20), title, font=title_font, fill=title_color)
    if body:
        lines = body.split("\n")
        y = y1 + 67
        for line in lines:
            draw.text((x1 + 24, y), line, font=body_font, fill=MUTED)
            y += 31


def arrow(draw, start, end, color=BLUE, width=5):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    for offset in (2.55, -2.55):
        point = (
            end[0] + length * math.cos(angle + offset),
            end[1] + length * math.sin(angle + offset),
        )
        draw.line([end, point], fill=color, width=width)


def diagram_canvas(title: str, subtitle: str):
    image = Image.new("RGB", (1600, 900), WHITE)
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 1600, 104), fill=NAVY)
    draw.text((56, 24), title, font=font(37, True), fill=WHITE)
    draw.text((56, 117), subtitle, font=font(22), fill=MUTED)
    return image, draw


def create_architecture_diagram(path: Path) -> None:
    image, draw = diagram_canvas(
        "DeepGuard system architecture",
        "Live biometric signals and deepfake inference run in the browser; the server stores metadata only.",
    )
    rounded_box(draw, (70, 230, 365, 470), PALE_BLUE, BLUE, "User device", "Camera frames\nChallenge instructions\nReact interface")
    rounded_box(draw, (460, 195, 860, 505), PALE_TEAL, TEAL, "On-device AI", "MediaPipe Face Landmarker\nLiveness state machines\nTransformers.js deepfake model\nExplainable decision policy")
    rounded_box(draw, (970, 230, 1260, 470), LIGHT, DARK_BLUE, "FastAPI", "Input validation\nSession and review APIs\nMetrics aggregation")
    rounded_box(draw, (1320, 230, 1530, 470), PALE_GREEN, GREEN, "SQLite", "Numeric results\nChallenge status\nNo camera frames")
    arrow(draw, (365, 350), (460, 350), TEAL)
    arrow(draw, (860, 350), (970, 350), BLUE)
    arrow(draw, (1260, 350), (1320, 350), GREEN)
    draw.text((382, 309), "frames", font=font(18, True), fill=TEAL)
    draw.text((868, 309), "JSON metadata", font=font(18, True), fill=BLUE)
    draw.text((1266, 309), "SQL", font=font(18, True), fill=GREEN)
    rounded_box(
        draw,
        (160, 590, 1440, 775),
        "#F7F9FB",
        LINE,
        "Resource model",
        "No GPU inference server is required.\n"
        "Heavy inference is delegated to WebGPU or browser WASM.\n"
        "FastAPI and SQLite handle only the lightweight metadata workload.",
        NAVY,
    )
    image.save(path, quality=95)


def create_verification_diagram(path: Path) -> None:
    image, draw = diagram_canvas(
        "Verification lifecycle",
        "The workflow stabilizes the face, proves ordered motion, samples three frames, and applies a transparent policy.",
    )
    steps = [
        ("1", "Prepare", "Request camera\nDetect source type", PALE_BLUE, BLUE),
        ("2", "Blink", "Neutral -> close\n-> reopen", PALE_TEAL, TEAL),
        ("3", "Head", "Right -> left\n-> center -> depth", PALE_AMBER, AMBER),
        ("4", "Smile", "Neutral -> smile\n-> release", PALE_GREEN, GREEN),
        ("5", "Analyze", "3 face crops\nModel + quality", LIGHT, DARK_BLUE),
        ("6", "Decide", "Genuine / fake\n/ manual review", PALE_RED, RED),
    ]
    x = 55
    boxes = []
    for number, title, body, fill, outline in steps:
        box = (x, 255, x + 225, 500)
        boxes.append(box)
        draw.rounded_rectangle(box, radius=16, fill=fill, outline=outline, width=3)
        draw.ellipse((x + 18, 277, x + 68, 327), fill=outline)
        number_box = draw.textbbox((0, 0), number, font=font(24, True))
        tw = number_box[2] - number_box[0]
        draw.text((x + 43 - tw / 2, 286), number, font=font(24, True), fill=WHITE)
        draw.text((x + 82, 281), title, font=font(28, True), fill=NAVY)
        for idx, line in enumerate(body.split("\n")):
            draw.text((x + 22, 365 + idx * 34), line, font=font(19), fill=MUTED)
        x += 258
    for left, right in zip(boxes, boxes[1:]):
        arrow(draw, (left[2], 377), (right[0], 377), BLUE, 4)
    draw.rounded_rectangle((160, 610, 1440, 765), radius=16, fill="#F7F9FB", outline=LINE, width=2)
    draw.text((195, 637), "Completion gate", font=font(27, True), fill=NAVY)
    draw.text(
        (195, 687),
        "Each challenge advances only after stable evidence across multiple frames; poor face quality resets the current tracker.",
        font=font(21),
        fill=MUTED,
    )
    image.save(path, quality=95)


def create_liveness_diagram(path: Path) -> None:
    image, draw = diagram_canvas(
        "Liveness state machines",
        "Temporal ordering makes a replay or single still image less likely to pass than one-frame threshold checks.",
    )
    generic = [
        ("Neutral", PALE_BLUE, BLUE),
        ("Action", PALE_TEAL, TEAL),
        ("Release", PALE_GREEN, GREEN),
        ("Passed", PALE_GREEN, GREEN),
    ]
    x = 120
    for index, (label, fill, outline) in enumerate(generic):
        box = (x, 220, x + 265, 340)
        rounded_box(draw, box, fill, outline, label)
        if index < len(generic) - 1:
            arrow(draw, (x + 265, 280), (x + 335, 280), outline)
        x += 335
    draw.text((120, 385), "Blink and smile", font=font(25, True), fill=NAVY)
    draw.text(
        (340, 386),
        "3 neutral frames -> action frames -> 2 release frames",
        font=font(22),
        fill=MUTED,
    )
    head = [
        ("Center", "4 frames"),
        ("Right", "2 frames"),
        ("Left", "2 frames"),
        ("Recenter", "3 frames"),
        ("Depth", "3 frames"),
    ]
    x = 80
    y = 540
    for index, (label, frames) in enumerate(head):
        fill = [PALE_BLUE, PALE_TEAL, PALE_AMBER, PALE_BLUE, PALE_GREEN][index]
        outline = [BLUE, TEAL, AMBER, BLUE, GREEN][index]
        box = (x, y, x + 250, y + 130)
        draw.rounded_rectangle(box, radius=15, fill=fill, outline=outline, width=3)
        draw.text((x + 24, y + 22), label, font=font(27, True), fill=NAVY)
        draw.text((x + 24, y + 72), frames, font=font(20), fill=MUTED)
        if index < len(head) - 1:
            arrow(draw, (x + 250, y + 65), (x + 300, y + 65), outline, 4)
        x += 300
    draw.text((80, 705), "Head sequence", font=font(25, True), fill=NAVY)
    draw.text(
        (300, 706),
        "Yaw thresholds prove direction; a 10% face-scale change proves movement toward or away from the camera.",
        font=font(20),
        fill=MUTED,
    )
    image.save(path, quality=95)


def create_backend_diagram(path: Path) -> None:
    image, draw = diagram_canvas(
        "Backend API and persistence",
        "A small metadata API supports history, review status, and evaluation summaries without retaining biometric imagery.",
    )
    endpoints = [
        ("GET", "/api/health", "Availability and runtime health"),
        ("POST", "/api/sessions", "Create a validated verification record"),
        ("GET", "/api/sessions", "List and filter recent records"),
        ("PATCH", "/api/sessions/{id}/review", "Clear or escalate manual review"),
        ("GET", "/api/metrics", "Summary counts and seven-day trend"),
    ]
    y = 195
    for method, route, purpose in endpoints:
        draw.rounded_rectangle((75, y, 930, y + 100), radius=13, fill=LIGHT, outline=LINE, width=2)
        method_fill = GREEN if method == "GET" else BLUE if method == "POST" else AMBER
        draw.rounded_rectangle((96, y + 23, 205, y + 77), radius=9, fill=method_fill)
        draw.text((116, y + 35), method, font=font(19, True), fill=WHITE)
        draw.text((235, y + 20), route, font=font(23, True), fill=NAVY)
        draw.text((235, y + 56), purpose, font=font(18), fill=MUTED)
        y += 116
    rounded_box(draw, (1050, 255, 1515, 590), PALE_GREEN, GREEN, "verification_sessions", "id + timestamp\ndecision + confidence\nrisk + liveness + quality\nsource + runtime + latency\nchallenge JSON + review state")
    arrow(draw, (930, 465), (1050, 425), GREEN)
    draw.rounded_rectangle((1040, 650, 1525, 790), radius=15, fill=PALE_RED, outline=RED, width=2)
    draw.text((1070, 674), "Privacy boundary", font=font(26, True), fill=RED)
    draw.text((1070, 724), "No raw camera frames are inserted.", font=font(20), fill=INK)
    image.save(path, quality=95)


def create_module_diagram(path: Path) -> None:
    image, draw = diagram_canvas(
        "Frontend module dependency map",
        "The verification view coordinates focused libraries; other views consume API and evaluation modules independently.",
    )
    rounded_box(draw, (610, 170, 990, 300), PALE_BLUE, BLUE, "App.tsx", "Navigation + lazy views")
    nodes = [
        ((80, 400, 405, 550), "VerificationView", PALE_TEAL, TEAL),
        ((455, 400, 780, 550), "SessionsView", LIGHT, DARK_BLUE),
        ((830, 400, 1155, 550), "EvaluationView", PALE_AMBER, AMBER),
        ((1205, 400, 1530, 550), "SystemView", PALE_GREEN, GREEN),
    ]
    for box, title, fill, outline in nodes:
        rounded_box(draw, box, fill, outline, title)
        arrow(draw, (800, 300), ((box[0] + box[2]) // 2, box[1]), outline, 4)
    libs = [
        ((60, 690, 310, 800), "face.ts", BLUE),
        ((345, 690, 595, 800), "liveness.ts", TEAL),
        ((630, 690, 880, 800), "deepfake.ts", DARK_BLUE),
        ((915, 690, 1165, 800), "decision.ts", RED),
        ((1200, 690, 1450, 800), "api.ts", GREEN),
    ]
    for box, title, outline in libs:
        draw.rounded_rectangle(box, radius=12, fill=WHITE, outline=outline, width=3)
        draw.text((box[0] + 24, box[1] + 34), title, font=font(25, True), fill=NAVY)
    for box, _, outline in libs[:4]:
        arrow(draw, (242, 550), ((box[0] + box[2]) // 2, box[1]), outline, 3)
    arrow(draw, (617, 550), (1325, 690), GREEN, 3)
    arrow(draw, (992, 550), (1325, 690), GREEN, 3)
    image.save(path, quality=95)


def create_diagrams() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = {
        "architecture": ASSET_DIR / "system_architecture.png",
        "verification": ASSET_DIR / "verification_lifecycle.png",
        "liveness": ASSET_DIR / "liveness_state_machines.png",
        "backend": ASSET_DIR / "backend_api_persistence.png",
        "modules": ASSET_DIR / "frontend_module_map.png",
    }
    create_architecture_diagram(diagrams["architecture"])
    create_verification_diagram(diagrams["verification"])
    create_liveness_diagram(diagrams["liveness"])
    create_backend_diagram(diagrams["backend"])
    create_module_diagram(diagrams["modules"])
    return diagrams


def line_count(path: Path) -> int:
    if not path.exists():
        return 0
    return len(path.read_text(encoding="utf-8", errors="replace").splitlines())


def all_source_paths() -> list[str]:
    return [item for _, items in SOURCE_GROUPS for item in items]


def add_code_excerpt(document: Document, relative_path: str, start: int, end: int) -> None:
    path = ROOT / relative_path
    lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
    selected = lines[start - 1 : end]
    p = document.add_paragraph(style="Code Listing")
    shade_paragraph(p, "F5F7FA", LINE)
    for idx, value in enumerate(selected, start=start):
        run = p.add_run(f"{idx:04d} | {value.expandtabs(4)}")
        set_run_font(run, name="Consolas", size=7.5, color="23272D")
        if idx < end:
            run.add_break()


def add_full_code_listing(document: Document, relative_path: str) -> None:
    path = ROOT / relative_path
    heading = document.add_paragraph(style="Code File")
    set_run_font(heading.add_run(relative_path), size=13, color=DARK_BLUE, bold=True)
    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(5)
    set_run_font(meta.add_run(PURPOSES[relative_path] + " "), size=9, color=MUTED, italic=True)
    set_run_font(meta.add_run(f"({line_count(path)} lines)"), size=9, color=MUTED, bold=True)

    lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
    if not lines:
        lines = [""]
    chunk_size = 45
    for offset in range(0, len(lines), chunk_size):
        chunk = lines[offset : offset + chunk_size]
        p = document.add_paragraph(style="Code Listing")
        shade_paragraph(p, "F5F7FA", LINE)
        for local_index, value in enumerate(chunk):
            line_number = offset + local_index + 1
            safe_value = value.expandtabs(4).replace("\x0b", " ")
            run = p.add_run(f"{line_number:04d} | {safe_value}")
            set_run_font(run, name="Consolas", size=7.5, color="23272D")
            if local_index < len(chunk) - 1:
                run.add_break()


def add_cover(document: Document) -> None:
    for _ in range(5):
        document.add_paragraph()
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    set_run_font(kicker.add_run("TECHNICAL ARCHITECTURE AND SOURCE REFERENCE"), size=10, color=TEAL, bold=True)

    title = document.add_paragraph(style="Title")
    title.add_run("DeepGuard")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run("Deepfake Detection and Liveness Verification System")
    subtitle_2 = document.add_paragraph()
    subtitle_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_2.paragraph_format.space_after = Pt(30)
    set_run_font(
        subtitle_2.add_run("Complete Code Architecture, Data Flow, Algorithms, APIs, Tests, and Full Source Listings"),
        size=12,
        color=MUTED,
        italic=True,
    )

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(4)
    rule.paragraph_format.space_after = Pt(28)
    set_paragraph_border(rule, TEAL, 10, 8)

    metadata = [
        ("Document type", "Graduation project technical manual"),
        ("Architecture", "Browser AI + FastAPI metadata service + SQLite"),
        ("Repository", "TurkeyAbdo/deepguard-graduation-project"),
        ("Prepared", date.today().strftime("%d %B %Y")),
        ("Coverage", "Complete first-party application, backend, tests, and project tooling"),
    ]
    table = document.add_table(rows=1, cols=2)
    set_repeat_table_header(table.rows[0])
    for index, value in enumerate(("Document profile", "Details")):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(value), size=9, color=WHITE, bold=True)
    for label, value in metadata:
        row = table.add_row()
        set_cell_shading(row.cells[0], PALE_BLUE)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(label), size=9, color=NAVY, bold=True)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(value), size=9, color=INK)
    set_table_geometry(table, [2700, 6660])
    set_table_borders(table)

    document.add_paragraph()
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(24)
    set_run_font(
        note.add_run("This manual documents the demonstrable prototype as implemented, including its known limitations."),
        size=9.5,
        color=MUTED,
        italic=True,
    )
    document.add_page_break()


def add_front_matter(document: Document, bullet_id: int) -> None:
    document.add_heading("Document Purpose", level=1)
    document.add_paragraph(
        "This manual explains how the DeepGuard graduation project is organized, how data moves through the system, "
        "how browser-side computer vision and deepfake inference reach a final decision, and how the lightweight "
        "backend stores results. It also includes the complete first-party source code used to build, test, demonstrate, "
        "and evaluate the prototype."
    )
    add_callout(
        document,
        "Scope statement",
        "The implementation is a research prototype and controlled demonstration system. It is not a certified "
        "biometric product and its controlled evaluation is evidence of prototype behavior, not proof of universal "
        "performance against every identity, camera, lighting condition, or future deepfake method.",
        PALE_AMBER,
        AMBER,
    )

    document.add_heading("Reading Guide", level=2)
    add_bullet(document, "Chapters 1-4 explain the architecture, module boundaries, and runtime flow.", bullet_id)
    add_bullet(document, "Chapters 5-8 explain signal extraction, liveness, deepfake inference, and the decision policy.", bullet_id)
    add_bullet(document, "Chapters 9-13 cover APIs, persistence, evaluation, testing, deployment, and extension points.", bullet_id)
    add_bullet(document, "Appendices A-E reproduce the complete first-party source listings with line numbers.", bullet_id)

    document.add_heading("Contents", level=1)
    contents = [
        "1. Executive Architecture Summary",
        "2. Repository and Module Structure",
        "3. End-to-End Verification Lifecycle",
        "4. Frontend Application Architecture",
        "5. Face Signal Extraction with MediaPipe",
        "6. Temporal Liveness State Machines",
        "7. Browser Deepfake Model",
        "8. Explainable Decision Policy",
        "9. Backend API and SQLite Persistence",
        "10. Sessions, Evaluation, and System Views",
        "11. Testing and Evidence Architecture",
        "12. Runtime, Resources, Privacy, and Deployment",
        "13. Maintenance and Extension Guide",
        "Appendices A-E. Complete Source Code",
    ]
    for item in contents:
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run(item), size=10.5, color=DARK_BLUE, bold=True)
    document.add_page_break()


def add_core_manual(document: Document, diagrams: dict[str, Path], bullet_id: int, number_id: int) -> None:
    document.add_heading("1. Executive Architecture Summary", level=1)
    document.add_paragraph(
        "DeepGuard is a local-first verification prototype. The browser owns the camera, extracts facial motion, "
        "executes a quantized image-classification model, and applies an explainable decision policy. The FastAPI "
        "service receives only derived session metadata and stores it in SQLite. This division keeps biometric imagery "
        "inside the user device and avoids a continuously provisioned GPU inference server."
    )
    add_picture(
        document,
        diagrams["architecture"],
        6.35,
        "DeepGuard system architecture",
        "Diagram showing the user camera feeding browser-side MediaPipe, liveness, deepfake inference, and decision logic; only JSON metadata proceeds to FastAPI and SQLite.",
        "Figure 1. DeepGuard local-first system architecture and privacy boundary.",
    )
    add_callout(
        document,
        "Main engineering decision",
        "Compute follows the camera. MediaPipe and Transformers.js run through WebGPU when available and browser "
        "WASM otherwise. The backend handles validation, records, review state, and metrics, which are inexpensive "
        "compared with server-side video inference.",
        PALE_TEAL,
        TEAL,
    )

    document.add_heading("1.1 Technology Stack", level=2)
    add_table(
        document,
        ["Layer", "Technology", "Responsibility"],
        [
            ["User interface", "React 19 + TypeScript 6 + Vite 8", "Views, workflow state, responsive controls, and result presentation"],
            ["Face signals", "MediaPipe Face Landmarker", "Landmarks, blendshapes, bounding box, quality, yaw, blink, and smile"],
            ["Deepfake score", "Transformers.js + ONNX quantized model", "Three-frame image classification in WebGPU or WASM"],
            ["Decision", "TypeScript threshold policy", "Genuine, fake, or manual review with explicit reasons"],
            ["API", "FastAPI + Pydantic", "Validation, sessions, reviews, health, and metrics"],
            ["Persistence", "SQLite", "Small structured verification records; no raw camera frames"],
            ["Testing", "Vitest + Pytest", "Policy, liveness, evaluation math, API, persistence, and review behavior"],
        ],
        [1800, 2520, 5040],
    )

    document.add_heading("1.2 Runtime Invariants", level=2)
    add_bullet(document, "The browser never uploads the sampled camera canvases to the API.", bullet_id)
    add_bullet(document, "A face-quality value below 0.42 resets the current liveness tracker.", bullet_id)
    add_bullet(document, "A virtual-camera label is treated as a replay-risk signal and produces a fake outcome.", bullet_id)
    add_bullet(document, "Unavailable inference or low quality produces manual review rather than a confident automatic result.", bullet_id)
    add_bullet(document, "Controlled evaluation reports coverage separately from accuracy so reviewed cases are not hidden.", bullet_id)

    document.add_heading("2. Repository and Module Structure", level=1)
    document.add_paragraph(
        "The repository separates the runnable application under deepguard/ from final report artifacts under outputs/ "
        "and repeatable document/evaluation tooling under work/. Within the application, browser libraries are small and "
        "focused, views own presentation and workflow coordination, and the Python backend remains a single compact service."
    )
    add_picture(
        document,
        diagrams["modules"],
        6.35,
        "Frontend module dependency map",
        "Diagram showing App.tsx above four lazy-loaded views, with VerificationView coordinating face, liveness, deepfake, decision, and API libraries.",
        "Figure 2. Frontend module boundaries and primary dependencies.",
    )

    manifest_rows = []
    for relative_path in all_source_paths():
        path = ROOT / relative_path
        layer = relative_path.split("/")[1] if "/" in relative_path else "repository"
        manifest_rows.append([relative_path, layer, str(line_count(path)), PURPOSES[relative_path]])
    add_table(
        document,
        ["File", "Area", "Lines", "Purpose"],
        manifest_rows,
        [3300, 1260, 660, 4140],
        center_columns=(2,),
    )
    add_callout(
        document,
        "Excluded generated or binary artifacts",
        "The complete listings intentionally omit node_modules, pnpm-lock.yaml, the MediaPipe .task model, WASM "
        "binaries, SQLite databases, generated images/videos, and generated report documents. Those files are "
        "dependencies or outputs rather than first-party source. Their locations and roles are still documented.",
        LIGHT,
        DARK_BLUE,
    )

    document.add_heading("2.1 Important Ownership Boundaries", level=2)
    add_bullet(document, "Views coordinate user-facing behavior; libraries implement reusable signal, policy, API, and metric logic.", bullet_id)
    add_bullet(document, "The backend does not reimplement liveness or model inference, preventing two conflicting decision engines.", bullet_id)
    add_bullet(document, "The evaluation module calculates metrics; the evaluation view only renders those values.", bullet_id)
    add_bullet(document, "Demo-video scripts generate synthetic, clearly labeled material and are not part of live verification.", bullet_id)
    add_bullet(document, "src/App.css has no import and is legacy Vite starter residue; production styling comes from src/index.css.", bullet_id)

    document.add_heading("3. End-to-End Verification Lifecycle", level=1)
    add_picture(
        document,
        diagrams["verification"],
        6.35,
        "Verification lifecycle",
        "Six-step workflow: prepare camera, blink, ordered head movement, smile, analyze three face crops, and decide.",
        "Figure 3. Verification phases from camera preparation to final decision.",
    )
    document.add_paragraph(
        "VerificationView.tsx is the orchestration boundary. It owns the finite UI phase, camera stream, MediaPipe "
        "landmarker, animation loop, challenge index, challenge results, sampled crops, quality history, inference "
        "result, persistence state, and cleanup. Focused helper modules perform the calculations so the view does not "
        "duplicate signal algorithms."
    )
    for step in (
        "Preparation: request an ideal 1280 x 720 camera stream, inspect the selected device label, initialize MediaPipe, and begin warming the deepfake model.",
        "Signal loop: approximately every 75 ms, process the current video frame and derive face presence, quality, yaw, blink, smile, and face scale.",
        "Challenge progression: update only the active temporal state machine; display stable stage instructions and progress to the user.",
        "Sampling: after each successful challenge, capture a centered 384 x 384 face crop. Three passed challenges produce three model samples.",
        "Analysis: average deepfake probability across the samples, compute mean quality, and compute liveness as passed challenges divided by total challenges.",
        "Decision and storage: apply the policy, display the decision and reason, then POST a structured record. An API failure leaves the local result visible and marks it unsaved.",
    ):
        add_numbered(document, step, number_id)

    document.add_heading("3.1 Verification Phase Model", level=2)
    add_table(
        document,
        ["Phase", "Meaning", "Primary exit condition"],
        [
            ["idle", "No active camera or verification", "User starts verification"],
            ["preparing", "Camera and models initialize", "Face processing becomes available"],
            ["challenge", "One of three liveness checks is active", "All challenge state machines pass"],
            ["analyzing", "Three sampled crops are scored", "Inference and policy finish"],
            ["complete", "Result and session identifier are displayed", "User starts a new verification"],
            ["error", "Camera/model/runtime failure blocks progress", "User retries or changes device conditions"],
        ],
        [1440, 4320, 3600],
    )

    document.add_heading("3.2 Source Integrity", level=2)
    document.add_paragraph(
        "The application classifies camera labels containing OBS, virtual, ManyCam, Snap Camera, NDI, or DroidCam as "
        "virtual. A physical camera label is classified as physical; a missing label is unknown. This is a defensive "
        "prototype control, not a cryptographically trustworthy device attestation. Browser labels may be absent before "
        "permission, and sophisticated virtual devices may use neutral names."
    )

    document.add_heading("4. Frontend Application Architecture", level=1)
    document.add_paragraph(
        "App.tsx is a small stateful shell with four lazy-loaded views. The navigation state selects Verify, Sessions, "
        "Evaluation, or System. React Suspense keeps initial loading small, while a refresh token lets a completed "
        "verification tell SessionsView to reload without introducing a global state framework."
    )
    add_table(
        document,
        ["View", "User task", "Main dependencies"],
        [
            ["VerificationView", "Run live or synthetic attack verification", "face, liveness, deepfake, decision, api"],
            ["SessionsView", "Inspect and review saved sessions", "api, StatusBadge, Recharts"],
            ["EvaluationView", "Inspect controlled cases and metrics", "evaluationCases, evaluation, Recharts"],
            ["SystemView", "Inspect runtime and resource architecture", "Browser capability APIs"],
        ],
        [1920, 3900, 3540],
    )
    document.add_heading("4.1 Interface State and Cleanup", level=2)
    add_bullet(document, "MediaStream tracks are stopped when verification resets or the view unmounts.", bullet_id)
    add_bullet(document, "The requestAnimationFrame loop is canceled before a new run begins.", bullet_id)
    add_bullet(document, "Mutable high-frequency tracking data lives in refs to avoid unnecessary React renders.", bullet_id)
    add_bullet(document, "Stable progress and instructions are derived from challenge stages, reducing scattered or rapidly changing text.", bullet_id)
    add_bullet(document, "The result remains visible when persistence fails, preserving the user-facing outcome while clearly indicating it was not saved.", bullet_id)

    document.add_heading("4.2 Responsive and Accessibility Structure", level=2)
    document.add_paragraph(
        "src/index.css defines a 236 px desktop sidebar, a two-column verification workspace, explicit semantic colors, "
        "stable camera geometry, horizontal table overflow, and mobile navigation below 760 px. The verification layout "
        "becomes single-column below 1060 px and the camera uses a portrait 4:5 ratio on phones. Reduced-motion users "
        "receive animation-free behavior. Controls use text labels with recognizable Lucide icons rather than ambiguous decoration."
    )

    document.add_heading("5. Face Signal Extraction with MediaPipe", level=1)
    document.add_paragraph(
        "src/lib/face.ts initializes MediaPipe Face Landmarker from local WASM and model assets. It requests a GPU delegate "
        "first and retries on CPU if initialization fails. Only one face is tracked because the verification experience "
        "is intentionally single-subject."
    )
    add_table(
        document,
        ["Derived signal", "How it is obtained", "How it is used"],
        [
            ["Face box", "Minimum and maximum landmark coordinates", "Centering, size, guide overlay, and crop region"],
            ["Yaw", "Relative nose-to-left/right-cheek distances", "Ordered right, left, and recenter stages"],
            ["Blink", "Maximum eyeBlinkLeft/eyeBlinkRight blendshape", "Neutral, closed-eye action, and reopening"],
            ["Smile", "Average mouthSmileLeft/mouthSmileRight blendshape", "Neutral, smile action, and release"],
            ["Face scale", "Normalized bounding-box width/height", "10% closer-or-away depth proof"],
            ["Quality", "Face size, centering, and landmark confidence factors", "Tracker reset, review gate, and session evidence"],
            ["Face crop", "Centered square canvas resized to 384 x 384", "Three-sample deepfake classification"],
        ],
        [1680, 3840, 3840],
    )
    add_callout(
        document,
        "Why quality matters",
        "A single low-quality frame should not advance an action. The tracker resets below 0.42, and the final policy "
        "routes mean quality below 0.45 to manual review. This prevents a poor image from being reported as a confident genuine result.",
        PALE_AMBER,
        AMBER,
    )
    document.add_heading("5.1 MediaPipe Initialization Excerpt", level=2)
    add_code_excerpt(document, "deepguard/src/lib/face.ts", 18, 48)

    document.add_heading("6. Temporal Liveness State Machines", level=1)
    add_picture(
        document,
        diagrams["liveness"],
        6.35,
        "Liveness state machines",
        "Diagram showing neutral-action-release-passed stages for blink and smile, plus ordered center-right-left-recenter-depth stages for head movement.",
        "Figure 4. Temporal liveness design and ordered head-motion proof.",
    )
    document.add_paragraph(
        "The liveness layer does not declare success from one threshold crossing. It requires temporal order and multiple "
        "stable frames. Blink and smile use a generic neutral-action-release tracker. Head motion uses a compound tracker "
        "that proves right turn, left turn, recentering, and a depth change in sequence."
    )
    add_table(
        document,
        ["Challenge", "Neutral/action thresholds", "Temporal evidence"],
        [
            ["Blink", "neutral <= 0.20; action >= 0.42", "3 neutral, 1 action, 2 release frames"],
            ["Smile", "neutral <= 0.20; action >= 0.42", "3 neutral, 2 action, 2 release frames"],
            ["Head direction", "|center| <= 0.18; right >= 0.30; left <= -0.30", "4 center, 2 right, 2 left, 3 recenter frames"],
            ["Head depth", "|yaw| <= 0.24 and face scale changes >= 10%", "3 stable depth frames after recenter"],
        ],
        [1680, 4020, 3660],
    )
    add_callout(
        document,
        "Security interpretation",
        "This sequence is stronger than a still-image check and useful for a classroom prototype. It is not a complete "
        "presentation-attack-detection certification. High-assurance deployment would add randomized prompts, optical "
        "flow, challenge timing, camera attestation, infrared/depth signals where available, and trained spoof classifiers.",
        PALE_RED,
        RED,
    )
    document.add_heading("6.1 Ordered Head-Motion Excerpt", level=2)
    add_code_excerpt(document, "deepguard/src/lib/liveness.ts", 75, 133)

    document.add_heading("7. Browser Deepfake Model", level=1)
    document.add_paragraph(
        "src/lib/deepfake.ts loads onnx-community/Deep-Fake-Detector-v2-Model-ONNX through Transformers.js. WebGPU uses "
        "a q4f16 quantized runtime; WASM uses q8. Loading is lazy and cached as a singleton so a verification run does not "
        "recreate the pipeline. Three face crops are classified sequentially and their deepfake probabilities are averaged."
    )
    add_table(
        document,
        ["Concern", "Implementation", "Effect"],
        [
            ["Server cost", "Inference runs in the browser", "No dedicated model-serving GPU is required"],
            ["Compatibility", "WebGPU first, WASM fallback", "Modern and older compatible browsers can participate"],
            ["Download size", "Quantized ONNX runtime", "Smaller model/runtime than full-precision server packages"],
            ["Noise reduction", "Average of three challenge samples", "Less dependence on one transitional frame"],
            ["Failure behavior", "Unavailable model routes to review", "The system avoids pretending a score exists"],
        ],
        [1740, 3840, 3780],
    )
    add_callout(
        document,
        "Model limitation",
        "The controlled trials produced texture-risk scores near 50.78% across both genuine and synthetic cases. "
        "Therefore the current demonstration should not claim that the texture model alone provides reliable "
        "general-purpose deepfake discrimination. The demonstrable outcome combines source integrity, temporal "
        "liveness, quality, and the model score.",
        PALE_AMBER,
        AMBER,
    )
    document.add_heading("7.1 Inference Excerpt", level=2)
    add_code_excerpt(document, "deepguard/src/lib/deepfake.ts", 1, 40)

    document.add_heading("8. Explainable Decision Policy", level=1)
    document.add_paragraph(
        "src/lib/decision.ts converts source integrity, model availability, face quality, liveness, and deepfake risk "
        "into a deterministic decision. The order of the checks matters: strong risk signals are handled before the "
        "genuine gates, and ambiguous conditions are sent to review."
    )
    add_table(
        document,
        ["Priority", "Condition", "Decision", "Reason"],
        [
            ["1", "Virtual camera detected", "fake", "Replay-risk source"],
            ["2", "Model unavailable or quality < 0.45", "review", "Insufficient evidence"],
            ["3", "Deepfake risk >= 0.72 or liveness < 0.45", "fake", "Strong fake or failed liveness"],
            ["4", "Physical + risk < 0.60 + liveness >= 0.80 + quality >= 0.55", "genuine", "Trusted source and strong evidence"],
            ["5", "Any source + risk < 0.40 + liveness >= 0.80 + quality >= 0.55", "genuine", "Very low risk and strong evidence"],
            ["6", "Everything else", "review", "Ambiguous evidence"],
        ],
        [780, 3900, 1260, 3420],
        center_columns=(0, 2),
    )
    document.add_heading("8.1 Why a Real Face Can Show 51% Risk", level=2)
    document.add_paragraph(
        "The displayed deepfake-risk percentage is the model's texture score, not the probability that the whole person "
        "is fake. A normal physical-camera attempt can therefore show roughly 51% texture risk while still passing all "
        "motion challenges. The final policy considers all signals. In the threshold band between confident genuine and "
        "strong fake, manual review is the intended result. The UI and report should preserve this distinction instead "
        "of relabeling the texture score as total confidence."
    )
    add_code_excerpt(document, "deepguard/src/lib/decision.ts", 1, 39)

    document.add_heading("9. Backend API and SQLite Persistence", level=1)
    add_picture(
        document,
        diagrams["backend"],
        6.35,
        "Backend API and persistence",
        "Five FastAPI endpoints connected to a verification_sessions SQLite table, with a privacy boundary excluding raw camera frames.",
        "Figure 5. FastAPI endpoints, SQLite record, and the metadata-only privacy boundary.",
    )
    document.add_paragraph(
        "server/app.py creates a FastAPI application that can be pointed at a custom database with DEEPGUARD_DB. "
        "Pydantic validates incoming session and review records. SQLite rows store numeric evidence, model/runtime "
        "metadata, challenge results as JSON, notes, and review status. When a production client build exists, the same "
        "service can also serve dist/ as static content."
    )
    add_table(
        document,
        ["Method", "Path", "Input", "Output"],
        [
            ["GET", "/api/health", "None", "Service health"],
            ["POST", "/api/sessions", "Validated SessionCreate JSON", "Stored session with generated DG identifier"],
            ["GET", "/api/sessions", "limit and optional decision filter", "Recent sessions"],
            ["PATCH", "/api/sessions/{id}/review", "Review status and notes", "Updated session"],
            ["GET", "/api/metrics", "None", "Counts and seven-day trend"],
        ],
        [900, 2760, 2940, 2760],
    )
    document.add_heading("9.1 Database Fields", level=2)
    add_bullet(document, "Identity and timing: id, timestamp, latency_ms.", bullet_id)
    add_bullet(document, "Decision evidence: decision, confidence, deepfake_risk, liveness, quality.", bullet_id)
    add_bullet(document, "Runtime provenance: source_integrity, source_label, model_version, runtime.", bullet_id)
    add_bullet(document, "Challenge evidence: challenge_results serialized as JSON.", bullet_id)
    add_bullet(document, "Human workflow: notes and review_status.", bullet_id)
    document.add_heading("9.2 API Creation Excerpt", level=2)
    add_code_excerpt(document, "deepguard/server/app.py", 93, 160)

    document.add_heading("10. Sessions, Evaluation, and System Views", level=1)
    document.add_heading("10.1 Session Review Workspace", level=2)
    document.add_paragraph(
        "SessionsView loads recent sessions and aggregate metrics in parallel. Users can filter genuine, fake, and "
        "review outcomes, expand a record to inspect runtime/quality/latency details, and clear or escalate review state. "
        "A Recharts area chart shows the seven-day verification trend."
    )
    document.add_heading("10.2 Controlled Evaluation", level=2)
    document.add_paragraph(
        "EvaluationView renders eleven controlled cases: physical-camera genuine attempts and OBS replay attacks. "
        "src/lib/evaluation.ts computes coverage, confusion-matrix values, automatic-decision accuracy, overall correctness "
        "including review, precision, recall, specificity, F1, and mean/median latency. Separating coverage from accuracy "
        "makes abstention visible."
    )
    add_table(
        document,
        ["Metric", "Definition", "Why it matters"],
        [
            ["Coverage", "Automatic decisions / total cases", "Shows how often the system avoids manual review"],
            ["Auto accuracy", "(TP + TN) / automatic decisions", "Measures correctness only when the system commits"],
            ["Overall correct", "(TP + TN + correct reviews) / total", "Represents the controlled workflow including review"],
            ["Precision", "TP / (TP + FP)", "How often fake predictions are correct"],
            ["Recall", "TP / (TP + FN)", "How many replay attacks are detected"],
            ["Specificity", "TN / (TN + FP)", "How many genuine attempts are accepted"],
            ["F1", "Harmonic mean of precision and recall", "Balances attack detection quality"],
        ],
        [1680, 3420, 4260],
    )
    document.add_heading("10.3 Resource Evidence", level=2)
    add_table(
        document,
        ["Measured item", "Controlled value", "Interpretation"],
        [
            ["FastAPI idle memory", "25.7 MiB", "Small metadata service footprint"],
            ["Mean metadata API response", "11.0 ms", "No server-side video/model inference in the request"],
            ["SQLite database", "32 KiB for 27 sessions", "Structured records are compact"],
            ["Static package", "59.5 MiB", "Local model/WASM assets dominate the package, not server memory"],
        ],
        [2520, 2160, 4680],
    )
    add_callout(
        document,
        "How to present the server argument",
        "The system does require client compute and model assets; it does not require many inference servers for this "
        "prototype architecture. A small API host can serve many metadata requests because each user's browser performs "
        "the expensive vision work. Larger production traffic would still require ordinary API scaling, monitoring, "
        "backups, and possibly managed storage.",
        PALE_GREEN,
        GREEN,
    )

    document.add_heading("11. Testing and Evidence Architecture", level=1)
    document.add_paragraph(
        "The automated suite targets the parts most likely to produce silent security regressions: thresholds, stage "
        "ordering, abstention, metric denominators, validation, persistence, and review transitions. Frontend tests use "
        "Vitest; backend tests use Pytest and a temporary SQLite database."
    )
    add_table(
        document,
        ["Suite", "Coverage"],
        [
            ["decision.test.ts", "Virtual source, strong risk, low liveness, physical genuine gate, and manual review"],
            ["liveness.test.ts", "Blink/smile temporal stages, quality resets, head order, recenter, and depth"],
            ["evaluation.test.ts", "Coverage, confusion matrix, accuracy, F1, and latency math"],
            ["server/test_app.py", "Health, empty metrics, session round trip, aggregate metrics, and review update"],
        ],
        [2400, 6960],
    )
    document.add_heading("11.1 Verification Commands", level=2)
    add_code_excerpt(document, "deepguard/package.json", 6, 18)
    document.add_paragraph(
        "The repository has been verified with frontend linting, fourteen frontend tests, a production build, and four "
        "backend tests. These checks validate implementation consistency, but camera/model behavior should also be "
        "demonstrated on the professor's target device because browser GPU support and camera conditions vary."
    )

    document.add_heading("12. Runtime, Resources, Privacy, and Deployment", level=1)
    document.add_heading("12.1 Development Runtime", level=2)
    for step in (
        "Install the JavaScript dependencies with pnpm and Python dependencies from requirements.txt.",
        "Run the launcher or start FastAPI on 127.0.0.1:8000 and Vite on 127.0.0.1:5173.",
        "Allow camera permission. Vite proxies /api to the FastAPI service during development.",
        "For a single-service production demonstration, build the frontend to dist/ and let FastAPI serve it.",
    ):
        add_numbered(document, step, number_id)
    document.add_heading("12.2 Resource Answer", level=2)
    add_bullet(document, "GPU server: not required by the implemented architecture.", bullet_id)
    add_bullet(document, "Client GPU: optional; WebGPU improves speed and browser WASM provides a fallback.", bullet_id)
    add_bullet(document, "API server: lightweight CPU service for JSON validation, queries, and metrics.", bullet_id)
    add_bullet(document, "Database: SQLite is suitable for the demonstrated single-node prototype; production concurrency can move to PostgreSQL without changing browser inference.", bullet_id)
    add_bullet(document, "Bandwidth: camera frames stay local; only compact result metadata is transmitted.", bullet_id)
    add_bullet(document, "Static assets: the model and WASM files must be downloaded/cached by each client, which shifts cost from server compute to static delivery.", bullet_id)

    document.add_heading("12.3 Privacy and Security Controls", level=2)
    add_bullet(document, "Raw frames are held transiently in browser video/canvas objects and are not persisted by the API.", bullet_id)
    add_bullet(document, "The database stores scores and challenge outcomes; treat them as sensitive biometric-adjacent metadata.", bullet_id)
    add_bullet(document, "CORS is limited to local development origins in the current prototype.", bullet_id)
    add_bullet(document, "A public deployment must add HTTPS, authentication, authorization, retention policy, audit logging, rate limiting, and encrypted backups.", bullet_id)
    add_bullet(document, "Model files and client code are public in the repository, so thresholds are inspectable; production anti-spoofing should not depend on secrecy.", bullet_id)

    document.add_heading("13. Maintenance and Extension Guide", level=1)
    document.add_heading("13.1 Safe Extension Points", level=2)
    add_table(
        document,
        ["Goal", "Primary files", "Recommended change"],
        [
            ["Add a challenge", "liveness.ts, VerificationView.tsx, liveness.test.ts", "Create a tracker, instruction stages, progress mapping, and tests"],
            ["Replace deepfake model", "deepfake.ts", "Preserve a normalized 0-1 score and explicit unavailable state"],
            ["Tune decision thresholds", "decision.ts, decision.test.ts", "Change policy and add boundary/evaluation evidence together"],
            ["Use PostgreSQL", "server/app.py", "Introduce a repository layer while preserving API contracts"],
            ["Add authentication", "server/app.py, api.ts", "Protect session/review routes and bind reviews to users"],
            ["Add production evaluation", "evaluationCases.ts, evaluation.ts", "Use labeled, diverse, ethically sourced datasets and report uncertainty"],
            ["Remove legacy styling", "App.css", "Delete only after confirming there are still no imports"],
        ],
        [1980, 3300, 4080],
    )
    document.add_heading("13.2 Recommended Production Roadmap", level=2)
    recommendations = [
        "Calibrate the deepfake score on a larger held-out dataset and display texture risk separately from total decision confidence.",
        "Randomize challenge order and parameters, enforce timing windows, and add optical-flow or landmark-trajectory features.",
        "Add a trained presentation-attack model covering screen replay, print, mask, and virtual-camera scenarios.",
        "Measure WebGPU/WASM latency, memory, and accuracy across target laptops and phones before defining device support.",
        "Add authenticated review roles, immutable audit events, retention controls, and encrypted production storage.",
        "Move to a managed database only when concurrent write volume, availability, or multi-instance deployment requires it.",
        "Adopt model/version monitoring so every session can be traced to thresholds, runtime, and model artifacts.",
    ]
    for item in recommendations:
        add_numbered(document, item, number_id)
    add_callout(
        document,
        "Final architecture assessment",
        "The project is fully demonstrable as a local-first graduation prototype. Its strongest implemented contribution "
        "is the integration of ordered liveness, source integrity, lightweight browser inference, explainable policy, "
        "review workflow, and evidence reporting. Production claims should remain proportional to dataset size and attack coverage.",
        PALE_BLUE,
        BLUE,
    )


def add_appendices(document: Document) -> None:
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Complete Source Code Appendices", level=1)
    document.add_paragraph(
        "The following appendices reproduce the complete first-party text source used by the application and its "
        "supporting evaluation/report workflow. Line numbers are added for reference and are not part of the repository files."
    )
    add_callout(
        document,
        "Listing convention",
        "Each code line is prefixed with a four-digit reference number. Generated dependencies, lockfiles, databases, "
        "binary models/WASM, images, videos, and DOCX outputs are documented but not reproduced as text.",
        LIGHT,
        DARK_BLUE,
    )
    for group_title, relative_paths in SOURCE_GROUPS:
        document.add_heading(group_title, level=1)
        for relative_path in relative_paths:
            add_full_code_listing(document, relative_path)


def set_headers_and_footers(document: Document) -> None:
    sections = list(document.sections)
    for index, section in enumerate(sections):
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        if index > 0:
            section.header.is_linked_to_previous = True
            section.footer.is_linked_to_previous = True

    header = sections[0].header
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("DEEPGUARD  |  CODE ARCHITECTURE"), size=8.5, color=MUTED, bold=True)
    set_paragraph_border(p, LINE, 4, 4)

    footer = sections[0].footer
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("Technical manual  |  Page "), size=8.5, color=MUTED)
    add_page_field(p)


def preset_audit(document: Document) -> None:
    section = document.sections[0]
    expected = {
        "top_margin": Inches(1),
        "bottom_margin": Inches(1),
        "left_margin": Inches(1),
        "right_margin": Inches(1),
    }
    for name, value in expected.items():
        actual = getattr(section, name)
        if abs(actual - value) > 2:
            raise AssertionError(f"{name} mismatch: {actual} vs {value}")
    if document.styles["Normal"].font.name != "Calibri":
        raise AssertionError("Normal font is not Calibri")
    for table in document.tables:
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        if tbl_w is None or tbl_w.get(qn("w:w")) != str(CONTENT_DXA):
            raise AssertionError("A table is missing 9360 DXA width")
        tbl_ind = table._tbl.tblPr.find(qn("w:tblInd"))
        if tbl_ind is None or tbl_ind.get(qn("w:w")) != str(TABLE_INDENT_DXA):
            raise AssertionError("A table is missing 120 DXA indent")


def main() -> None:
    missing = [relative for relative in all_source_paths() if not (ROOT / relative).exists()]
    if missing:
        raise FileNotFoundError(f"Missing source files: {missing}")
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = create_diagrams()

    document = Document()
    document.core_properties.title = "DeepGuard Full Code Architecture Guide"
    document.core_properties.subject = "Deepfake detection and liveness verification technical architecture"
    document.core_properties.author = "DeepGuard Graduation Project"
    document.core_properties.keywords = "DeepGuard, deepfake, liveness, React, MediaPipe, FastAPI, SQLite, source code"
    configure_styles(document)
    bullet_id, number_id = create_numbering(document)

    add_cover(document)
    add_front_matter(document, bullet_id)
    add_core_manual(document, diagrams, bullet_id, number_id)
    add_appendices(document)
    set_headers_and_footers(document)
    preset_audit(document)

    document.save(OUTPUT)
    reopened = Document(OUTPUT)
    print(f"Created: {OUTPUT}")
    print(f"Paragraphs: {len(reopened.paragraphs)}")
    print(f"Tables: {len(reopened.tables)}")
    print(f"Inline shapes: {len(reopened.inline_shapes)}")
    print(f"Source files: {len(all_source_paths())}")
    print(f"Source lines: {sum(line_count(ROOT / path) for path in all_source_paths())}")


if __name__ == "__main__":
    main()
