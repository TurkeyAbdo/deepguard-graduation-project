from __future__ import annotations

import copy
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCUMENT_SKILL = Path(
    r"C:\Users\Lenovo\.codex\plugins\cache\openai-primary-runtime\documents\26.715.12143\skills\documents"
)
sys.path.insert(0, str(DOCUMENT_SKILL / "scripts"))
from table_geometry import apply_table_geometry, column_widths_from_weights, section_content_width_dxa  # noqa: E402


SOURCE = ROOT / "outputs" / "FULL_REPORT_COMPLETE_WITH_CHAPTER_5_BW_WORKFLOWS_UI.docx"
OUTPUT = ROOT / "outputs" / "FULL_REPORT_FINAL_WITH_EVALUATION.docx"
ASSETS = ROOT / "work" / "evaluation_assets"


def set_bidi(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:bidi")) is None:
        p_pr.append(OxmlElement("w:bidi"))


def format_run(run, size: float = 12, bold: bool = False, color: str = "181A1D") -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._r.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), "Arial")
    if r_pr.find(qn("w:rtl")) is None:
        r_pr.append(OxmlElement("w:rtl"))


def insert_paragraph(document: Document, anchor, *, text: str = "", style=None):
    paragraph = document.add_paragraph(style=style)
    anchor._p.addprevious(paragraph._p)
    if text:
        run = paragraph.add_run(text)
        format_run(run)
    set_bidi(paragraph)
    return paragraph


def insert_heading(document: Document, anchor, text: str, *, major: bool = False):
    paragraph = insert_paragraph(document, anchor, style=anchor.style)
    run = paragraph.add_run(text)
    format_run(run, 14 if major else 12.5, True)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(14 if major else 10)
    paragraph.paragraph_format.space_after = Pt(6 if major else 4)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = major
    return paragraph


def insert_body(document: Document, anchor, text: str, *, lead: str | None = None, bold: bool = False):
    paragraph = insert_paragraph(document, anchor)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.5
    if lead:
        lead_run = paragraph.add_run(lead)
        format_run(lead_run, 12, True)
        body_run = paragraph.add_run(text)
        format_run(body_run, 12)
    else:
        run = paragraph.add_run(text)
        format_run(run, 12, bold)
    return paragraph


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = tc_pr.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
        tc_pr.append(shade)
    shade.set(qn("w:fill"), fill)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def format_cell(cell, text: str, *, header: bool = False, center: bool = False) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    set_bidi(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    format_run(run, 10.5 if header else 10.25, header, "FFFFFF" if header else "24272B")
    if header:
        shade_cell(cell, "24272B")


def insert_table(document: Document, anchor, headers: list[str], rows: list[list[str]], weights: list[float]):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, value in enumerate(headers):
        format_cell(table.rows[0].cells[index], value, header=True, center=True)
    set_repeat_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            format_cell(cells[index], value, center=index == 1)
            if row_index % 2 == 1:
                shade_cell(cells[index], "F3F4F5")
    bidi_visual = OxmlElement("w:bidiVisual")
    table._tbl.tblPr.append(bidi_visual)
    total_width = section_content_width_dxa(document.sections[0])
    widths = column_widths_from_weights(weights, total_width)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=total_width,
        indent_dxa=130,
        cell_margins_dxa={"top": 100, "bottom": 100, "start": 130, "end": 130},
    )
    anchor._p.addprevious(table._tbl)
    spacer = insert_paragraph(document, anchor)
    spacer.paragraph_format.space_after = Pt(2)
    return table


def insert_figure(document: Document, anchor, image_path: Path, caption: str, width: float = 6.1):
    picture_paragraph = insert_paragraph(document, anchor)
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.space_before = Pt(6)
    picture_paragraph.paragraph_format.space_after = Pt(3)
    picture_paragraph.paragraph_format.keep_with_next = True
    run = picture_paragraph.add_run()
    picture = run.add_picture(str(image_path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption)

    caption_paragraph = insert_paragraph(document, anchor)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(8)
    caption_run = caption_paragraph.add_run(caption)
    format_run(caption_run, 10, True, "50575E")
    return picture_paragraph


def replace_text(document: Document, old: str, new: str) -> None:
    paragraph = next(item for item in document.paragraphs if item.text.strip() == old)
    paragraph.text = ""
    set_bidi(paragraph)
    run = paragraph.add_run(new)
    size = 14 if old.startswith(("5.4 ", "5.5 ", "5.6 ")) else 12
    format_run(run, size, old.startswith(("5.4 ", "5.5 ", "5.6 ")))


def main() -> None:
    document = Document(SOURCE)
    anchor = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "5.4 التوصيات")

    intro = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip().startswith("يعرض هذا الفصل النتائج"))
    intro.text = ""
    set_bidi(intro)
    intro_run = intro.add_run(
        "يعرض هذا الفصل النتائج التي توصل إليها المشروع، ويقدم تقييمًا تجريبيًا مضبوطًا للأداء وزمن الاستجابة واستهلاك الموارد، ثم يوضح قابلية التشغيل بموارد محدودة وتوصيات التطوير المستقبلية والمصادر والخاتمة العامة للمشروع."
    )
    format_run(intro_run)

    old_result = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip().startswith("أظهر النموذج الأولي قدرة واعدة"))
    old_result.text = ""
    set_bidi(old_result)
    result_run = old_result.add_run(
        "أظهرت مجموعة الاختبار المضبوطة المكونة من 11 تجربة كاميرا أن النظام أصدر قرارًا آليًا في 10 حالات، بينما حوّل حالة حقيقية منخفضة الجودة إلى المراجعة البشرية. وتُعرض المقاييس التفصيلية وحدودها في القسم 5.4."
    )
    format_run(result_run)

    method_note = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip().startswith("ملاحظة منهجية: لا ينبغي"))
    method_note.text = ""
    set_bidi(method_note)
    note_lead = method_note.add_run("ملاحظة منهجية: ")
    format_run(note_lead, 12, True)
    note_body = method_note.add_run(
        "جميع النسب الواردة أدناه خاصة بعينة تجريبية صغيرة ومضبوطة، ولا تمثل دقة النموذج على مجموعات بيانات عامة أو على مستخدمين وأجهزة وظروف تصوير متنوعة."
    )
    format_run(note_body)

    insert_heading(document, anchor, "5.4 التقييم التجريبي وقياس الأداء", major=True)
    insert_body(
        document,
        anchor,
        "أُجري التقييم على التجارب ذات الحقيقة المرجعية المعروفة فقط، بهدف قياس المسار الكامل للنظام من التقاط الوجه إلى إصدار القرار. اعتُبرت هجمات إعادة العرض عبر كاميرا OBS الافتراضية فئة موجبة (Attack)، واعتُبرت جلسات الكاميرا المادية لمستخدم حقيقي فئة سالبة (Genuine).",
    )

    insert_heading(document, anchor, "5.4.1 نطاق العينة ومنهجية القياس")
    insert_table(
        document,
        anchor,
        ["نوع التجربة", "العدد", "النتيجة المسجلة"],
        [
            ["كاميرا مادية - حقيقة مرجعية: حقيقي", "8", "7 حقيقي + 1 مراجعة بشرية"],
            ["إعادة عرض عبر OBS - حقيقة مرجعية: هجوم", "3", "3 مخاطر مرتفعة"],
            ["جلسات غير موسومة", "10", "مستبعدة من حساب الدقة"],
            ["محاكاة واجهة ثابتة", "6", "مستبعدة لأنها لا تنفذ استدلالًا فعليًا"],
        ],
        [2.4, 0.9, 3.0],
    )
    insert_body(
        document,
        anchor,
        "تُحسب الدقة وPrecision وRecall وSpecificity وF1 على الحالات التي أصدر فيها النظام قرارًا آليًا فقط، بينما تُعرض Coverage بصورة مستقلة لبيان نسبة الحالات التي لم تُحوّل إلى المراجعة البشرية. ويمنع هذا الفصل بين الدقة والتغطية إخفاء أثر حالات الامتناع عن القرار.",
    )

    insert_heading(document, anchor, "5.4.2 نتائج التصنيف")
    insert_table(
        document,
        anchor,
        ["المقياس", "القيمة", "التفسير"],
        [
            ["حجم العينة الموسومة", "11", "8 حالات حقيقية و3 هجمات إعادة عرض"],
            ["التغطية الآلية", "90.9%", "10 قرارات آلية من أصل 11"],
            ["دقة الحالات المغطاة", "100.0%", "10 قرارات صحيحة من أصل 10 قرارات آلية"],
            ["Precision للهجوم", "100.0%", "لم تُسجل حالة حقيقية مصنفة كهجوم"],
            ["Recall للهجوم", "100.0%", "كُشفت هجمات OBS الثلاث"],
            ["Specificity", "100.0%", "قُبلت 7 حالات حقيقية ولم تُرفض أي حالة حقيقية آليًا"],
            ["F1-score", "100.0%", "على القرارات الآلية داخل العينة المضبوطة"],
            ["الصحة الكلية مع المراجعة", "90.9%", "10 نتائج صحيحة وحالة واحدة للمراجعة"],
        ],
        [1.8, 1.0, 3.5],
    )
    insert_figure(
        document,
        anchor,
        ASSETS / "confusion_matrix.png",
        "الشكل (5-1): مصفوفة الالتباس لنتائج الاختبار المضبوط مع إظهار قرار المراجعة البشرية.",
    )
    insert_body(
        document,
        anchor,
        "يجب تفسير النسب السابقة بحذر؛ فهجمات OBS كُشفت أساسًا بإشارة سلامة المصدر الافتراضي، كما أن عدد التجارب صغير وينتمي إلى بيئة تشغيل واحدة. لذلك تُعد هذه النتائج إثباتًا لسلامة سير العمل وليست تقديرًا نهائيًا لقدرة نموذج كشف التزييف على التعميم.",
        lead="قيد التفسير: ",
    )

    insert_heading(document, anchor, "5.4.3 زمن الاستجابة")
    insert_body(
        document,
        anchor,
        "بلغ متوسط الزمن من بدء التحقق حتى إصدار القرار 14.65 ثانية، وبلغ الوسيط 14.73 ثانية، وتراوح الزمن بين 8.53 و22.67 ثانية. يشمل القياس تنفيذ تحديات الحيوية والتقاط العينات والاستدلال بعد توفر النموذج في ذاكرة التخزين المؤقت للمتصفح، ولا يشمل زمن تنزيل النموذج في أول تشغيل.",
    )
    insert_figure(
        document,
        anchor,
        ASSETS / "latency_by_trial.png",
        "الشكل (5-2): زمن التحقق الكامل لكل تجربة في العينة المضبوطة.",
    )

    insert_heading(document, anchor, "5.4.4 قياس الموارد وإثبات عدم الحاجة إلى خادم GPU")
    insert_body(
        document,
        anchor,
        "أظهر القياس المحلي أن خدمة FastAPI استهلكت 25.7 MiB من الذاكرة في حالة الخمول، وأن متوسط الاستجابة لخمسين طلبًا محليًا إلى واجهة بيانات المقاييس بلغ 11.0 ms. وبلغ حجم قاعدة SQLite التي تحتوي على 27 سجل جلسة 32 KiB، بينما بلغ حجم الحزمة الساكنة للمتصفح 59.5 MiB، ومعظمها ملفات WASM ونموذج معالم الوجه التي تُنزّل وتُخزّن مؤقتًا على جهاز المستخدم.",
    )
    insert_figure(
        document,
        anchor,
        ASSETS / "resource_profile.png",
        "الشكل (5-3): ملف الموارد المقاس للنموذج الأولي المحلي.",
    )
    insert_body(
        document,
        anchor,
        "تثبت هذه البنية أن التشغيل الأساسي لا يحتاج إلى خادم GPU؛ فالاستدلال العصبي وتحليل الإطارات يحدثان داخل المتصفح باستخدام WebGPU أو WASM، بينما يستقبل الخادم بيانات JSON رقمية صغيرة فقط. يلزم خادم أقوى عند التدريب أو عند زيادة عدد المستخدمين المتزامنين وخدمات الإدارة، وليس لتنفيذ جلسة العرض الفردية.",
    )

    insert_heading(document, anchor, "5.4.5 حدود التقييم")
    limitations = [
        ("حجم العينة: ", "تتكون العينة من 11 تجربة فقط ولا تشمل تنوعًا كافيًا في الأشخاص والأجهزة والإضاءة والخلفيات."),
        ("معايرة المصنف: ", "استقر خطر النسيج عند نحو 50.8% في جميع التجارب المقاسة، مما يدل على عدم كفاية معايرة نموذج النسيج الحالي، واعتماد القرار التجريبي بدرجة أكبر على الحيوية وسلامة المصدر."),
        ("سلامة المصدر: ", "اكتشاف OBS يعتمد على اسم جهاز الكاميرا الافتراضية، وهو إجراء دفاعي أولي يمكن التحايل عليه إذا تغير اسم الجهاز؛ لذلك لا يعد بديلًا كاملًا لتقنيات Presentation Attack Detection."),
        ("نطاق الهجوم: ", "لم تُختبر بعد هجمات الصور المطبوعة والشاشات الخارجية والأقنعة ثلاثية الأبعاد والإضاءة الضعيفة والحجب الجزئي."),
        ("الاعتماد على المتصفح: ", "يتغير الأداء باختلاف دعم WebGPU وسرعة الجهاز، وقد ينتقل النظام إلى WASM على الأجهزة غير المدعومة."),
    ]
    for lead, text in limitations:
        insert_body(document, anchor, text, lead=lead)

    insert_heading(document, anchor, "5.4.6 توصيات التحقق العلمي التالي")
    recommendations = [
        ("مجموعة بيانات مستقلة: ", "اختبار النموذج على FaceForensics++ أو DFDC أو بيانات محلية موسومة لم تدخل في التدريب، مع تقسيم ثابت للتدريب والتحقق والاختبار."),
        ("معايرة العتبة: ", "حساب منحنى ROC وAUC واختيار العتبة على مجموعة التحقق، ثم تثبيتها قبل اختبار المجموعة النهائية."),
        ("مقاييس القياسات الحيوية: ", "إضافة FAR وFRR وEER ومقاييس APCER وBPCER وفق مبادئ ISO/IEC 30107-3."),
        ("اختبار متعدد الأجهزة: ", "إعادة القياس على أجهزة مختلفة وفي ظروف إضاءة ومسافات وجودة كاميرا متنوعة، مع فصل زمن أول تشغيل عن زمن التشغيل بعد التخزين المؤقت."),
        ("سجل تقييم مستقل: ", "إضافة حقل للحقيقة المرجعية وبيانات السيناريو في قاعدة تقييم منفصلة عن سجل التشغيل الفعلي، لضمان إمكانية إعادة حساب النتائج ومراجعتها."),
    ]
    for lead, text in recommendations:
        insert_body(document, anchor, text, lead=lead)

    replace_text(document, "5.4 التوصيات", "5.5 التوصيات")
    replace_text(document, "5.5 المصادر والمراجع", "5.6 المصادر والمراجع")
    replace_text(document, "5.6 الخاتمة", "5.7 الخاتمة")

    core = document.core_properties
    core.title = "Deepfake Detection and Liveness Detection - Final Evaluated Report"
    core.subject = "Graduation project report with controlled evaluation and measured resource profile"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
